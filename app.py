import streamlit as st
import pandas as pd
import io
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

st.set_page_config(page_title="Mithila Audit System - Unicode", layout="wide")

# ==============================================================================
# ०. रोमनलाई नेपाली युनिकोडमा बदल्ने फ्री फङ्ग्सन (GOOGLE INPUT TOOLS STYLE)
# ==============================================================================
def convert_to_nepali(text_input):
    if not text_input or text_input.strip() == "":
        return text_input
    
    # यदि प्रयोगकर्ताले अङ्ग्रेजी मात्र लेख्न खोजेको होइन भने मात्र रूपान्तरण गर्ने
    try:
        url = f"https://google.com{text_input}&itc=ne-t-i0-und&num=1"
        response = requests.get(url, timeout=5)
        if response.status_code == 200:
            data = response.json()
            if data[0] == "SUCCESS":
                return data[1][0][1][0]
    except:
        pass  # इन्टरनेट नभएमा वा एरर आएमा सादा टेक्स्ट नै फर्काउने
    return text_input

# ==============================================================================
# १. प्रयोगकर्ता इनपुट संकलन खण्ड (SIDEBAR)
# ==============================================================================
st.sidebar.title("अडिट प्रतिवेदन सेटिङहरू")

st.sidebar.header("पाना १: कभर पेज विवरण")
# इनपुट कोठाहरूमा रोमन टाइप गरेपछि अन्त्यमा रूपान्तरण बटन राखिएको छ
shakha_input = st.sidebar.text_input("शाखा कार्यालय (रोमनमा टाइप गर्नुहोस):", "choharwa, siraha")
shakha_name = convert_to_nepali(shakha_input)
st.sidebar.caption(f"📝 रूपान्तरण: {shakha_name}")

pesh_input = st.sidebar.text_input("पेस गरिएको मिति:", "31 baishakh 2083")
pesh_miti = convert_to_nepali(pesh_input)

aud1_input = st.sidebar.text_input("आन्तरिक लेखापरीक्षक १:", "roshan gurung")
auditor_1 = convert_to_nepali(aud1_input)

aud2_input = st.sidebar.text_input("आन्तरिक लेखापरीक्षक २:", "suresh patel")
auditor_2 = convert_to_nepali(aud2_input)

st.sidebar.header("पाना २: संक्षिप्त जानकारी मितिहरू")
audit_start_date = st.sidebar.text_input("लेखापरिक्षण सुरु मिति (देखि):", "२०८२/०४/०१")
audit_end_date = st.sidebar.text_input("लेखापरिक्षण अन्तिम मिति (सम्म):", "२०८२/११/३०")
audit_duration_input = st.sidebar.text_input("लागेको समय:", "3 din")
audit_duration_days = convert_to_nepali(audit_duration_input)

audit_period_input = st.sidebar.text_input("लेखापरिक्षण कार्य अवधि:", "2083/01/07 dekhi 2083/01/09 samma")
audit_period_str = convert_to_nepali(audit_period_input)

# ==============================================================================
# मुख्य स्क्रिन लेआउट (MAIN INTERFACE)
# ==============================================================================
tabs = st.tabs(["📄 पाना १ (Cover Page)", "📊 पाना २ (Summary & Cash Audit)"])

# ------------------------------------------------------------------------------
# TAB 1: COVER PAGE DESIGN
# ------------------------------------------------------------------------------
with tabs:
    st.markdown("<p style='text-align: right; font-style: italic; color: gray;'>Mithila Laghubitta Bittiya Sanstha Ltd / Choharwa Branch</p>", unsafe_allow_html=True)
    st.write("---")
    st.markdown("<h1 style='text-align: center; font-size: 42px; margin-top: 50px;'>आन्तरिक लेखापरीक्षण प्रतिबेदन</h1>", unsafe_allow_html=True)
    st.markdown(f"<h3 style='text-align: center; font-size: 26px;'>शाखा कार्यालय {shakha_name}</h3>", unsafe_allow_html=True)
    st.markdown("<h1 style='text-align: center; margin-top: 40px; color: gray;'>↕</h1>", unsafe_allow_html=True)
    st.markdown("<h4 style='text-align: center; font-size: 20px; margin-top: 40px;'><b>पेश गरिएको :</b></h4>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; font-size: 18px;'>मिथिला लघुवित्त वित्तीय संस्था लिमिटेड<br>केन्द्रीय कार्यालय, ढल्केवर, धनुषा ।</p>", unsafe_allow_html=True)
    st.markdown("<h1 style='text-align: center; margin-top: 30px; color: gray;'>↕</h1>", unsafe_allow_html=True)
    st.markdown("<h4 style='text-align: center; font-size: 20px; margin-top: 30px;'><b>प्रतिबेदन पेश गर्ने :</b></h4>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; font-size: 18px;'>आन्तरिक लेखापरीक्षण बिभाग<br>केन्द्रीय कार्यालय, ढल्केवर, धनुषा ।</p>", unsafe_allow_html=True)
    st.markdown(f"<p style='text-align: center; font-size: 16px;'><b>आ. लेखापरीक्षकहरु :</b> {auditor_1} / {auditor_2}</p>", unsafe_allow_html=True)
    st.write("---")
    st.markdown("<table style='width:100%'><tr><td><small>Mithila Laghubitta Bittiya Sanstha Ltd / Internal Audit Department</small></td><td style='text-align:right'><small>Page 1</small></td></tr></table>", unsafe_allow_html=True)

# ------------------------------------------------------------------------------
# TAB 2: PAGE 2 DESIGN WITH TRANSCRIPTION FOR DYNAMIC ENTRIES
# ------------------------------------------------------------------------------
with tabs:
    st.markdown("<p style='text-align: right; font-style: italic; color: gray;'>Mithila Laghubitta Bittiya Sanstha Ltd / Choharwa Branch</p>", unsafe_allow_html=True)
    st.markdown(f"<h3 style='text-align: right;'>{pesh_miti}</h3>", unsafe_allow_html=True)
    
    st.header("१. शाखाको संक्षिप्त जानकारी")
    
    col_ind1, col_ind2, col_ind3 = st.columns(3)
    with col_ind1:
        staff_in = st.text_input("शाखामा कार्यरत कर्मचारी:", "4 jana, (1 jana messenger sahita)")
        staff_count = convert_to_nepali(staff_in)
        center_in = st.text_input("केन्द्र संख्या (#):", "90 wota")
        center_count = convert_to_nepali(center_in)
        member_in = st.text_input("सदस्य संख्या (#):", "1212")
        member_count = convert_to_nepali(member_in)
    with col_ind2:
        borrower_in = st.text_input("ऋणी सदस्य संख्या (#):", "884")
        borrower_count = convert_to_nepali(borrower_in)
        total_loan_in = st.text_input("लगानीमा रहीरहेको रकम (रु.):", "12,01,81,920.92/-")
        total_loan = convert_to_nepali(total_loan_in)
        total_saving_in = st.text_input("बचत परिचालन (रु.):", "3,04,१४,487.54/-")
        total_saving = convert_to_nepali(total_saving_in)
    with col_ind3:
        npa_loan_in = st.text_input("भाखा नाघेको कर्जा (रु.):", "2,11,00,123/- (175 jana)")
        npa_loan = convert_to_nepali(npa_loan_in)
        npa_percent_in = st.text_input("भाखा नाघेको कर्जाको प्रतिशत:", "17.55%")
        npa_percent = convert_to_nepali(npa_percent_in)
        watch_list_in = st.text_input("सुक्ष्म निगरानीमा रहेको कर्जा (*):", "3,75,21,996/- (308 jana)")
        watch_list = convert_to_nepali(watch_list_in)
        productivity_in = st.text_input("कर्मचारी उत्पादकत्व:", "606 jana / 45 wota")
        productivity = convert_to_nepali(productivity_in)

    indicator_data = {
        "क्र.स": [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14],
        "सूचकहरु": [
            "शाखा कार्यालयको नाम :", "कार्यालयको ठेगाना :", "लेखापरिक्षण गरिएको अवधि :", 
            "लेखापरिक्षणमा लागेको समयः", "शाखामा कार्यरत कर्मचारी :", "केन्द्र संख्या :", 
            "सदस्य संख्या :", "ऋणी सदस्य संख्या :", "लगानीमा रहीरहेको रकम :", 
            "बचत परिचालन :", "भाखा नाघेको कर्जा :", "भाखा नाघेको कर्जाको प्रतिशत :", 
            "सुक्ष्म निगरानीमा रहेको कर्जा :", "कर्मचारी उत्पादकत्व / प्रति कर्मचारी केन्द्र :"
        ],
        "विवरण": [
            shakha_name, shakha_name, f"{audit_start_date} देखी {audit_end_date} सम्म",
            f"{audit_period_str} सम्म {audit_duration_days}", staff_count, center_count,
            member_count, borrower_count, total_loan, total_saving, npa_loan, npa_percent, watch_list, productivity
        ]
    }
    st.table(pd.DataFrame(indicator_data).set_index("क्र.स"))
    
    note_star_in = st.text_input("फुटनोट (*):", "NPA ma naaeko tara bhakha naghera...")
    note_star = convert_to_nepali(note_star_in)
    note_hash_obj = st.text_input("फुटनोट (#):", f"# miti {pesh_miti} samma ko")
    note_hash = convert_to_nepali(note_hash_obj)
    
    st.header("२. आन्तरीक लेखापरिक्षण गर्दा अपनाईएका विधिहरु तथा आधारहरु")
    st.write("क) कार्यालयको लेखा सम्बन्धका विषयबस्तुहरुको पुर्ण लेखा परिक्षण तथा कार्याक्रम तथा अन्य विषयबस्तुहरुको नमून परिक्षण विधिबाट गरिएको छ ।")
    st.write("ख) शाखा कार्यालयले लेखा परिक्षण अवधिभरमा उपलब्ध गराएको तथ्याङ्कको आधारमा यो प्रतिवेदन तयार पारिएको छ ।")
    st.write(f"ग) मिति {audit_end_date.split('/') if '/' in audit_end_date else audit_end_date} मसान्तको सन्तुलन परिक्षणको आधारमा लेखा सम्बन्धी कारोबारहरुको परिक्षण गरिएको छ ।")

    # --------------------------------------------------------------------------
    # ३. नगद तथा ढुकुटीको निरिक्षण (EXCEL FORMAT)
    # --------------------------------------------------------------------------
    st.header("३. नगद तथा ढुकुटीको निरिक्षण")
    
    notes = [1000, 500, 100, 50, 20, 10, 5, 2, 1]
    noteCounts = {}
    totalCash = 0.0

    col_form1, col_form2 = st.columns(2)
    with col_form1:
        st.write("#### दर र संख्या प्रविष्टि (Denomination Quantity)")
        for n in notes:
            noteCounts[n] = st.number_input(f"रु. {n} को संख्या:", min_value=0, value=0, step=1, key=f"notes_qty_{n}")
            totalCash += n * noteCounts[n]
            
    with col_form2:
        st.write("#### प्रणाली र निरिक्षण विवरण")
        system_cash = st.number_input("सफ्टवेयर (CBS) मा देखिएको नगद मौज्दात (रु.):", min_value=0.0, value=0.0, step=1.0)
        inspection_date_in = st.text_input("भौतिक निरीक्षण गरिएको मिति/समय:", "miti 2083/01/06")
        inspection_date = convert_to_nepali(inspection_date_in)

    farak_amount = totalCash - system_cash
    if farak_amount == 0:
        remarks_str = "मिलेको"
    elif farak_amount < 0:
        remarks_str = f"रु. {abs(farak_amount):,}/- ले ढुकुटीमा कमी"
    else:
        remarks_str = f"रु. {farak_amount:,}/- ले ढुकुटीमा बढी"

    cash_dino_rows = []
    for n in notes:
        cash_dino_rows.append({
            "cash dino": f"रु. {n}",
            "Quantity": f"{noteCounts[n]:,}",
            "Amount": f"रु. {n * noteCounts[n]:,}/-"
        })
    df_dino = pd.DataFrame(cash_dino_rows)
    
    st.write("##### क) नगद बर्गिकरण तालिका")
    st.table(df_dino)

    st.write("##### ख) ढुकुटीको नगद भिडान तालिका")
    comparison_rows = [
        {"विवरण": "भौतिक निरीक्षणमा पाइएको कुल नगद मौज्दात (Total Cash)", "मौज्दात रकम": f"रु. {totalCash:,}/-", "कैफियत": inspection_date},
        {"विवरण": "सफट्वेयर (CBS) मा देखिएको नगद मौज्दात", "मौज्दात रकम": f"रु. {system_cash:,}/-", "कैफियत": ""},
        {"विवरण": "फरक / रेकन्सिल रकम (Difference Amount)", "मौज्दात रकम": f"रु. {farak_amount:,}/-", "कैफियत": remarks_str}
    ]
    st.table(pd.DataFrame(comparison_rows).set_index("विवरण"))

    # --------------------------------------------------------------------------
    # ३.१ Day End तथा दैनिक कारोबार
    # --------------------------------------------------------------------------
    st.header("३.१ Day End तथा दैनिक कारोबार")
    st.write("निरीक्षण अवधिमा जम्मा २४१ दिन मध्ये शाखाले अधिकांश दिनहरुमा समय मै अर्थात् कारोबार भएको दिन Day End नगरेको पाइयो । शाखाको ढिलो Day End गरेको विवरण यस प्रकार रहेको छ ।")
    
    col_de1, col_de2 = st.columns(2)
    with col_de1:
        de_ontime_in = st.text_input("समयमा Day End गरेको जम्मा दिन संख्या:", "9 din")
        de_ontime = convert_to_nepali(de_ontime_in)
        de_delay1_in = st.text_input("१ दिन देखि ४ दिन सम्म ढिला Day End गरेको जम्मा:", "106 din")
        de_delay1 = convert_to_nepali(de_delay1_in)
    with col_de2:
        de_delay2_in = st.text_input("५ दिन देखि ९ दिन सम्म ढिला Day End गरेको जम्मा:", "80 din")
        de_delay2 = convert_to_nepali(de_delay2_in)
        de_delay3_in = st.text_input("१० दिन देखि १४ दिन सम्म ढिला Day End गरेको जम्मा:", "46 din")
        de_delay3 = convert_to_nepali(de_delay3_in)

    st.write("---")
    st.markdown("<table style='width:100%'><tr><td><small>Mithila Laghubitta Bittiya Sanstha Ltd / Internal Audit Department</small></td><td style='text-align:right'><small>Page 2</small></td></tr></table>", unsafe_allow_html=True)

# ==============================================================================
# WORD (.DOCX) COMPILER
# ==============================================================================
def build_word_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    p_top = doc.add_paragraph()
    p_top.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_top.add_run("Mithila Laghubitta Bittiya Sanstha Ltd / Choharwa Branch").italic = True
    
    doc.add_paragraph("\n\n")
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.add_run("आन्तरिक लेखापरीक्षण प्रतिबेदन\n").font.size = Pt(28)
    p_title.add_run(f"शाखा कार्यालय {shakha_name}").font.size = Pt(18)
    
    doc.add_paragraph("\n\n\n")
    p_pesh = doc.add_paragraph()
    p_pesh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pesh.add_run("पेश गरिएको :\n").bold = True
    p_pesh.add_run("मिथिला लघुवित्त वित्तीय संस्था लिमिटेड\nकेन्द्रीय कार्यालय, ढल्केवर, धनुषा ।\n")
    
    doc.add_paragraph("\n\n")
    p_done = doc.add_paragraph()
    p_done.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_done.add_run("प्रतिबेदन पेश गर्ने :\n").bold = True
    p_done.add_run("आन्तरिक लेखापरीक्षण विभाग\nकेन्द्रीय कार्यालय, ढल्केवर, धनुषा ।\n")
    p_done.add_run(f"\nआ. लेखापरीक्षकहरु : {auditor_1} / {auditor_2}").bold = True
    
    doc.add_page_break()
    
    p2_top = doc.add_paragraph()
    p2_top.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2_top.add_run(f"{pesh_miti}\n").bold = True
    
    doc.add_heading("१. शाखाको संक्षिप्त जानकारी", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows.cells
    hdr_cells.text, hdr_cells.text, hdr_cells.text = 'क्र.स', 'सूचकहरु', 'विवरण'
    hdr_cells._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="E6E6E6"/>'))
    
    for idx in range(len(indicator_data["क्र.स"])):
        row_cells = table.add_row().cells
        row_cells.text = str(indicator_data["क्र.स"][idx])
        row_cells.text = indicator_data["सूचकहरु"][idx]
        row_cells.text = str(indicator_data["विवरण"][idx])
        
    doc.add_paragraph(note_star)
    doc.add_paragraph(note_hash)
    
    doc.add_heading("२. आन्तरीक लेखापरिक्षण गर्दा अपनाईएका विधिहरु तथा आधारहरु", level=1)
    doc.add_paragraph("क) कार्यालयको लेखा सम्बन्धका विषयबस्तुहरुको पुर्ण लेखा परिक्षण तथा कार्याक्रम तथा अन्य विषयबस्तुहरुको नमून परिक्षण विधिबाट गरिएको छ ।")
    doc.add_paragraph("ख) शाखा कार्यालयले लेखा परिक्षण अवधिभरमा उपलब्ध गराएको तथ्याङ्कको आधारमा यो प्रतिवेदन तयार पारिएको छ ।")
    doc.add_paragraph(f"ग) मिति {audit_end_date} मसान्तको सन्तुलन परिक्षणको आधारमा लेखा सम्बन्धी कारोबारहरुको परिक्षण गरिएको छ ।")
    
    doc.add_heading("३. नगद तथा ढुकुटीको निरिक्षण", level=1)
    table_dino = doc.add_table(rows=1, cols=3)
    table_dino.style = 'Table Grid'
    hd = table_dino.rows.cells
    hd.text, hd.text, hd.text = 'cash dino', 'Quantity', 'Amount'
    
    for n in notes:
        r_cells = table_dino.add_row().cells
        r_cells.text = f"रु. {n}"
        r_cells.text = f"{noteCounts[n]:,}"
        r_cells.text = f"रु. {n * noteCounts[n]:,}/-"
        
    doc.add_paragraph("\n")
    table_comp = doc.add_table(rows=1, cols=3)
    table_comp.style = 'Table Grid'
    tc = table_comp.rows.cells
    tc.text, tc.text, tc.text = 'विवरण', 'मौज्दात रकम', 'कैफियत'
    
    for row in comparison_rows:
        rc = table_comp.add_row().cells
        rc.text = row["विवरण"]
        rc.text = row["मौज्दात रकम"]
        rc.text = row["कैफियत"]
        
    doc.add_heading("३.१ Day End तथा दैनिक कारोबार", level=1)
    doc.add_paragraph("निरीक्षण अवधिमा जम्मा २४१ दिन मध्ये शाखाले अधिकांश दिनहरुमा समय मै अर्थात् कारोबार भएको दिन Day End नगरेको पाइयो । शाखाको ढिलो Day End गरेको विवरण यस प्रकार रहेको छ ।")
    doc.add_paragraph(f"- समयमा Day End गरेको जम्मा दिन संख्या: {de_ontime}")
    doc.add_paragraph(f"- १ दिन देखि ४ दिन सम्म ढिला Day End गरेको जम्मा: {de_delay1}")
    doc.add_paragraph(f"- ५ दिन देखि ९ दिन सम्म ढिला Day End गरेको जम्मा: {de_delay2}")
    doc.add_paragraph(f"- १० दिन देखि १४ दिन सम्म ढिला Day End गरेको जम्मा: {de_delay3}")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ==============================================================================
# DOWNLOAD BUTTONS
# ==============================================================================
st.sidebar.markdown("---")
st.sidebar.subheader("प्रतिवेदन डाउनलोड गर्नुहोस्")
st.sidebar.download_button(
    label="📝 Word (.docx) फाइल डाउनलोड गर्नुहोस्",
    data=build_word_document(),
    file_name=f"Audit_Report_Mithila_{shakha_name}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
